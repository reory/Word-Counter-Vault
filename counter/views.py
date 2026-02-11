import matplotlib
matplotlib.use("Agg") # Forces matplotlib to use no GUI or Tkinter. causes issues.
import matplotlib.pyplot as plt

from django.shortcuts import render
from django.contrib import messages
from django.http import HttpResponse
from django.shortcuts import redirect
from django.urls import reverse
from django.template.loader import render_to_string
from django.templatetags.static import static

from docx import Document
from docx.shared import Inches

from weasyprint import HTML
from typing import Any
from io import BytesIO

from .services.text_extractors import(
    extract_text_from_txt,
    extract_text_from_pdf,
    extract_text_from_docx,
)
from .services.analysis import get_word_frequencies
from .services.summarizer import generate_summary
from .services.quality_insights import analyze_quality

def counter(request):
    context: dict[str, Any] = {"on": "active"}
    
    text = request.session.pop("analysis_text", "")

    if request.method == "POST":
        
        # Text from text area.
        text = request.POST.get("texttocount", "")
        # Text from file upload.
        if "file" in request.FILES:
            uploaded_file = request.FILES["file"]
            filename = uploaded_file.name.lower()
            
            if filename.endswith(".txt"):
                text = extract_text_from_txt(uploaded_file)
            
            elif filename.endswith("pdf"):
                text = extract_text_from_pdf(uploaded_file)

            elif filename.endswith(".docx"):
                text = extract_text_from_docx(uploaded_file)

        # Store text in session and redirect.
        request.session["analysis_text"] = text
        return redirect("counter:home")
        
        # Analysis metrics.
    if text:
        # Run analysis.
        word_count = len(text.split())
        char_count = len(text)
        sentence_count = text.count(".") + text.count("!") + text.count("?")
        paragraph_count = len([p for p in text.split("\n") if p.strip()])
        reading_time = max(1, round(word_count / 200))
        
        # Context block.
        context["text"] = text
        context["word_count"] = word_count
        context["char_count"] = char_count
        context["sentence_count"] = sentence_count
        context["paragraph_count"] = paragraph_count
        context["reading_time"] = reading_time
        context["has_result"] = True
        
        # Summary data from summarizer.py
        summary_data = generate_summary(text)
        
        context["summary"] = summary_data["summary"]
        context["bullets"] = summary_data["bullets"]
        context["topics"] = summary_data["topics"]
        
        # Summary of data on quality of text.
        quality = analyze_quality(text)

        context.update({
            "longest_sentence": quality["longest_sentence"],
            "ttr": quality["ttr"],
            "overused": quality["overused"],
            "passive_count": quality["passive_count"]
        })
        # Save everything to session for a PDF export.
        request.session["analysis_text"] = text 
        request.session["summary"] = summary_data["summary"] 
        request.session["bullets"] = summary_data["bullets"] 
        request.session["topics"] = summary_data["topics"] 
        request.session["word_count"] = word_count 
        request.session["longest_sentence"] = quality["longest_sentence"] 
        request.session["ttr"] = quality["ttr"] 
        request.session["overused"] = quality["overused"] 
        request.session["passive_count"] = quality["passive_count"]

        # Message for the user.
        messages.success(request, "File uploaded and analyzed successfully!")

    return render(request, "counter/counter.html", context)

def word_frequency_chart(request):

    text = request.GET.get("text", "")

    if not text.strip():
        return HttpResponse("No text provided", status=400)
    
    # Get top 10 words.
    freq = get_word_frequencies(text)
    words = [w for w, _ in freq]
    counts = [c for _, c in freq]
        
    # Create the chart.
    plt.figure(figsize=(8, 4))
    plt.bar(words, counts, color="#2D2D35FF")
    plt.title("Top 10 Most Common Words")
    plt.xlabel("Words")
    plt.ylabel("Frequency")
    plt.tight_layout()
   
    
    # Save to memory.
    buffer = BytesIO()
    plt.savefig(buffer, format="png")
    plt.close()
    buffer.seek(0)

    return HttpResponse(buffer.getvalue(), content_type="image/png")

from io import BytesIO
import matplotlib.pyplot as plt
from docx.shared import Inches

def generate_word_frequency_chart_image(text):
    """This is for the Word docx chart."""
    freq = get_word_frequencies(text)
    words = [w for w, _ in freq]
    counts = [c for _, c in freq]

    plt.figure(figsize=(8, 4))
    plt.bar(words, counts, color="#2D2D35FF")
    plt.title("Top 10 Most Common Words")
    plt.xticks(rotation=45)
    plt.tight_layout()

    buffer = BytesIO()
    plt.savefig(buffer, format="png")
    plt.close()
    buffer.seek(0)

    return buffer

def export_pdf(request):

    logo_url = request.build_absolute_uri(static("counter/img/python_developer.png"))
    
    #Collect the data you want to include in the PDF.
    # Pull everything from the session.
    context = {
        "text": request.session.get("analysis_text"),
        "summary": request.session.get("summary"),
        "bullets": request.session.get("bullets"),
        "topics": request.session.get("topics"),
        "word_count": request.session.get("word_count"),
        "longest_sentence": request.session.get("longest_sentence"),
        "ttr": request.session.get("ttr"),
        "overused": request.session.get("overused"),
        "passive_count": request.session.get("passive_count"),
        "chart_url": request.build_absolute_uri(
            reverse("counter:chart") 
            + f"?text={request.session.get('analysis_text')}"
        ),
        "logo_url": logo_url,
    }
    
    # Render the HTML template with the context.
    html_string = render_to_string("export/pdf_template.html", context)
    # Convert HTML to PDF.
    html = HTML(string=html_string)
    pdf_bytes = html.write_pdf()
    # Return as a downloadable file.
    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response ["Content-Disposition"] = "attachment; filename=analysis_report.pdf"
    return response

def export_docx(request):

    # Pull everything for the session.
    text = request.session.get("analysis_text")
    summary = request.session.get("summary") 
    bullets = request.session.get("bullets") 
    topics = request.session.get("topics") 
    word_count = request.session.get("word_count") 
    longest_sentence = request.session.get("longest_sentence") 
    ttr = request.session.get("ttr") 
    overused = request.session.get("overused") 
    passive_count = request.session.get("passive_count")

    # Create the word document.
    doc = Document()

    # Title.
    doc.add_heading("Word Counter Report", level=1)

    # Word count.
    doc.add_heading("Word Count", level=2)
    doc.add_paragraph(str(word_count))

    # Summary.
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary)

    # Bullets.
    doc.add_heading("Key Points", level=2)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # Word frequency chart
    doc.add_heading("Word Frequency Chart", level=2)
    chart_image = generate_word_frequency_chart_image(text)
    doc.add_picture(chart_image, width=Inches(6))
    
    #Center the image.
    doc.paragraphs[-1].alignment = 1

    # Topics.
    doc.add_heading("Detected Topics", level=2)
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    # Quality Insights.
    doc.add_heading("Quality Insights", level=2)
    doc.add_paragraph(f"Longest Sentence: {longest_sentence}")
    doc.add_paragraph(f"Type-Token Ratio: {ttr}") 
    doc.add_heading("Overused Words", level=3)
    for word, count in overused:
        doc.add_paragraph(f"{word} - {count} times", style="List Bullet") 
    doc.add_paragraph(f"Passive Voice Count: {passive_count}")
    
    # Original Text.
    doc.add_heading("Original Text", level=2) 
    doc.add_paragraph(text) 
    
    # Save to memory.
    from io import BytesIO 
    buffer = BytesIO() 
    doc.save(buffer) 
    buffer.seek(0) 
    
    # Return as download. 
    response = HttpResponse(
        buffer.getvalue(), 
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" ) 
    response["Content-Disposition"] = "attachment; filename=analysis_report.docx" 
    return response