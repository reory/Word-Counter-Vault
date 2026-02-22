import matplotlib
matplotlib.use("Agg") # Forces matplotlib to use no GUI or Tkinter.
import matplotlib.pyplot as plt
from django.template.loader import render_to_string
from weasyprint import HTML
import base64
from io import BytesIO
from .analysis import get_word_frequencies
from docx import Document
from docx.shared import Inches
from . import word_selectors

def get_export_data(user, pk=None, session=None):
    """
    Unified helper to gather data for either PDF or Word.
    It checks the Vault (PK) first, then falls back to the Session.
    """
    # Prevents 'NoneType' errors if the function is called without a session
    if session is None:
        # Initialize as empty dict so 
        # .get() calls safely return None instead of crashing
        session = {}

    if pk:
        record = word_selectors.get_record_for_user(user, pk)
        return {
            "text": record.original_text,
            "summary": record.summary,
            "topics": record.topics,
            "word_count": record.word_count,
            "bullets": record.bullets,
            "longest_sentence": record.longest_sentence,
            "ttr": record.ttr,
            "overused": record.overused,
            "passive_count": record.passive_count,
        }
    
    # Fallback to session data
    return {
        "text": session.get("analysis_text"),
        "summary": session.get("summary"),
        "bullets": session.get("bullets"),
        "topics": session.get("topics"),
        "word_count": session.get("word_count"),
        "longest_sentence": session.get("longest_sentence"),
        "ttr": session.get("ttr"),
        "overused": session.get("overused"),
        "passive_count": session.get("passive_count"),
    }

def generate_pdf_report(context):
    """Handles the logic of turning a data dictionary into a PDF binary."""

    # Handle the chart generation internally.
    text = context.get('text', '')

    if text:
        freq = get_word_frequencies(text)
        words = [w for w, _ in freq]
        counts = [c for _, c in freq]

        plt.figure(figsize=(8, 4))
        plt.bar(words, counts, color='#2d2d35ff')
        plt.title('Top 10 Most Common Words')
        plt.xticks(rotation=45)
        plt.tight_layout()

        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        plt.close()

        image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        context['chart_data'] = f'data:image/png;base64,{image_base64}'

    # Render the HTML to PDF.
    html_string = render_to_string('export/pdf_template.html', context)
    html = HTML(string=html_string)
    return html.write_pdf()

def generate_word_frequency_chart_image(text):
    """Internal helper for Word doc chart generation."""
    freq = get_word_frequencies(text)
    words = [w for w, _ in freq]
    counts = [c for _, c in freq]

    plt.figure(figsize=(8, 4))
    plt.bar(words, counts, color="#2D2D35FF")
    plt.title("Top 10 Most Common Words")
    plt.xticks(rotation=45)
    plt.tight_layout()

    buffer = BytesIO()
    plt.savefig(buffer, format="png")
    plt.close()
    buffer.seek(0)
    return buffer

def generate_docx_report(data):
    """Constructs the full Word document from a data dictionary."""

    doc = Document()
    doc.add_heading('Word Counter Report', level=1)

    # Metrics
    doc.add_heading('Word Count', level=2)
    doc.add_paragraph(str(data.get('word_count', 0)))

    # Summary
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(data.get('summary', ''))

    doc.add_heading('Key Points', level=2)
    for b in data.get('bullets', []):
        doc.add_paragraph(b, style='List Bullet')

    # Chart
    doc.add_heading('Word Frequency Chart', level=2)
    chart_image = generate_word_frequency_chart_image(data.get('text', ''))
    doc.add_picture(chart_image, width=Inches(6))
    doc.paragraphs[-1].alignment = 1

    # Quality Insights
    doc.add_heading('Quality Insights', level=2)
    doc.add_paragraph(f"Longest Sentence: {data.get('longest_sentence', '')}")
    doc.add_paragraph(f"Type-Token Ratio: {data.get('ttr', 0)}")

    doc.add_heading('Overused Words', level=3)
    overused = data.get('overused', [])
    if overused:
        for word, count in overused:
            doc.add_paragraph(f"{word.capitalize()} - {count} times", style="List Bullet")
    
    else:
        doc.add_paragraph("No significant overused detected.")
    
    doc.add_paragraph(f"Passive Voice Count: {data.get('passive_count', 0)}")

    # Original text
    doc.add_heading('Original Text', level=1)
    doc.add_paragraph(data.get('text', ''))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()